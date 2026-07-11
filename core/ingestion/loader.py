from pathlib import Path
from typing import List
from langchain_core.documents import Document
import fitz
import hashlib
import re

def get_file_hash(file_path: str) -> str:
    with open(file_path, "rb") as f:
        return hashlib.md5(f.read()).hexdigest()


def clean_text(text: str) -> str:
    """
    Strip invisible/broken control characters left over from PDF or PPT
    conversion (things like \\x00, \\x0e, etc.), without touching normal
    letters, numbers, punctuation, or line breaks.
    """
    # remove control characters (the invisible junk bytes)
    text = re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)

    # collapse leftover extra blank lines / extra spaces caused by the cleanup
    text = re.sub(r'\n{3,}', '\n\n', text)
    text = re.sub(r'[ \t]{2,}', ' ', text)

    return text.strip()


def load_pdf(file_path: str) -> List[Document]:
    docs = []
    file_hash = get_file_hash(file_path)
    with fitz.open(file_path) as pdf:
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = clean_text(page.get_text())
            if text:
                docs.append(Document(
                    page_content=text,
                    metadata={
                        "source": file_path,
                        "filename": Path(file_path).name,
                        "page": page_num + 1,
                        "type": "pdf",
                        "file_hash": file_hash
                    }
                ))
    return docs


def load_txt(file_path: str) -> List[Document]:
    file_hash = get_file_hash(file_path)
    with open(file_path, "r", encoding="utf-8") as f:
        text = clean_text(f.read())
    return [Document(
        page_content=text,
        metadata={
            "source": file_path,
            "filename": Path(file_path).name,
            "page": 1,
            "type": "txt",
            "file_hash": file_hash
        }
    )]


def load_docx(file_path: str) -> List[Document]:
    from docx import Document as DocxDocument
    file_hash = get_file_hash(file_path)
    doc = DocxDocument(file_path)
    text = clean_text("\n".join([para.text for para in doc.paragraphs if para.text.strip()]))
    return [Document(
        page_content=text,
        metadata={
            "source": file_path,
            "filename": Path(file_path).name,
            "page": 1,
            "type": "docx",
            "file_hash": file_hash
        }
    )]


LOADERS = {
    ".pdf": load_pdf,
    ".txt": load_txt,
    ".docx": load_docx
}


def load_document(file_path: str) -> List[Document]:
    path = Path(file_path)

    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")

    ext = path.suffix.lower()

    if ext not in LOADERS:
        raise ValueError(f"Unsupported file type: {ext}. Supported: {list(LOADERS.keys())}")

    try:
        return LOADERS[ext](file_path)
    except Exception as e:
        raise RuntimeError(f"Failed to load {path.name}: {e}")